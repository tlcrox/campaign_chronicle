#!/usr/bin/env python3
"""
storyboard.py - build a Word storyboard from a scene CSV + transcript + images.

Moved from scripts/merge.py (behavior-preserving). The only change is that the
local timestamp parser was replaced by the shared pipeline.common.timecode
(killing one of the duplicate timecode implementations).

Public API:
    generate_storyboard(csv_path, transcript_path, image_folder, output_docx)

CLI:
    python3 -m pipeline.merge.storyboard <scene_dir> <session_text>
"""

from typing import Any, Callable

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import pandas as pd
import re
import json
from pathlib import Path
import sys
from pipeline.config import get_config
from pipeline.common.scenes import iter_scene_images, parse_scene_name

# Shared timecode parser (handles [MM:SS.ss] and [HH:MM:SS.ss]).
from pipeline.common.timecode import timestamp_to_seconds

import logging
logger = logging.getLogger(__name__)


def parse_txt_transcript(transcript_path):
    """Parse old .txt format: [TIME] Speaker: Dialogue"""
    entries = []
    with open(transcript_path, 'r', encoding='utf-8') as f:
        for line in f:
            match = re.match(r'(\[\d+:[\d\.]+\]|\[\d+:\d+:[\d\.]+\])\s*(.*?):\s*(.*)', line)
            if match:
                ts_raw, speaker, dialogue = match.groups()
                current_secs = timestamp_to_seconds(ts_raw)
                entries.append({
                    'start': current_secs,
                    'end': current_secs,
                    'speaker': speaker,
                    'text': dialogue.strip(),
                    'confidence': 1.0
                })
    return entries

def parse_json_transcript(transcript_path):
    """Parse new JSON format from whisperx with segments, speaker roles, confidence"""
    entries = []
    last_known_speaker = 'Unknown'
    try:
        with open(transcript_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        if 'segments' in data:
            for segment in data['segments']:
                start = segment.get('start', 0)
                end = segment.get('end', 0)
                text = segment.get('text', '').strip()
                speaker = segment.get('speaker', None)
                confidence = segment.get('confidence', 1.0)

                if speaker is None or speaker == 'MISSING':
                    speaker = last_known_speaker
                else:
                    last_known_speaker = speaker

                entries.append({
                    'start': start,
                    'end': end,
                    'speaker': speaker,
                    'text': text,
                    'confidence': confidence
                })
        else:
            logger.warning("JSON format doesn't match expected whisperx structure with 'segments'")

    except json.JSONDecodeError as e:
        logger.error(f"Error parsing JSON: {e}")

    return entries

def detect_transcript_format(transcript_path):
    """Detect format based on file extension: .json or .txt"""
    path_str = str(transcript_path).lower()
    return path_str.endswith('.json')

def generate_storyboard(csv_path, transcript_path, image_folder, output_docx):
    cfg = get_config()

    doc = Document()
    doc.add_heading(cfg.document_title, 0)

    # Blank page (title on page 1, content starts on page 2)
    doc.add_page_break()

    is_json = detect_transcript_format(transcript_path)

    if is_json:
        logger.info(f"Parsing JSON transcript: {transcript_path}")
        entries = parse_json_transcript(transcript_path)
    else:
        logger.info(f"Parsing TXT transcript: {transcript_path}")
        entries = parse_txt_transcript(transcript_path)

    paragraph_count = 0

    logger.info(f"Processing {len(entries)} transcript entries")

    # Scenes are optional. The CSV is used ONLY to place scene images on the
    # timeline, so when there is no CSV or no image folder we skip image pairing
    # and emit a transcript-only document (e.g. audio-only / no-scenes sessions).
    image_entries = []  # (start_time, video, scene, image_path)
    df = None
    if csv_path and Path(csv_path).exists():
        # All CSVs have normalized headers on line 1 (PySceneDetect timing line
        # stripped by detect_scenes.sh).
        df = pd.read_csv(csv_path, header=0)
        logger.debug(f"Scene columns: {list(df.keys())} length {len(df)} in {image_folder}")

    if df is not None and image_folder and Path(image_folder).exists():
        # Pair each image to its CSV row by the (video, scene) key encoded in the
        # FILENAME (Scene-{video:02d}-{scene:03d}), NOT by position in the glob.
        # This survives the user curating the auto-generated set: a surviving
        # "Scene-02-008" is still matched to its own (video 2, scene 8) row. Older
        # single-key CSVs without a Video column fall back to Scene Number alone.
        scene_col = cfg.scene_number_column
        video_col = cfg.video_column
        has_video_col = video_col in df.columns
        for image_path in iter_scene_images(image_folder, "Scene-*"):
            parsed = parse_scene_name(image_path.name)
            if parsed is None:
                logger.warning(f"Skipping non-canonical image name: {image_path.name}")
                continue
            video_idx, scene_num = parsed
            if has_video_col:
                sel = df[(df[video_col] == video_idx) & (df[scene_col] == scene_num)]
            else:
                sel = df[df[scene_col] == scene_num]
            if len(sel):
                start_time = sel.iloc[0][cfg.start_time_column]
                image_entries.append((start_time, video_idx, scene_num, image_path))
            else:
                logger.warning(
                    f"Image {image_path.name} has no CSV row "
                    f"(video {video_idx}, scene {scene_num})"
                )
    else:
        logger.info("No scene CSV/images provided; generating a transcript-only document.")

    # Place images in timeline (start-time) order, breaking ties by (video, scene).
    image_entries.sort(key=lambda e: (e[0], e[1], e[2]))

    entry_idx = 0
    for image_start_time, video_idx, scene_num, image_path in image_entries:
        # Collect entries up to this image's start time
        entry_idx, added = _flush_speaker_entries(
            entries, entry_idx, doc, is_json, image_start_time)
        paragraph_count += added

        # Add image
        if paragraph_count > 0:
            doc.add_page_break()
        # Use Heading 2 for scene headers (for TOC)
        hdr = doc.add_heading(f"Scene {video_idx:02d}-{scene_num:03d}", level=2)
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(image_path), width=Inches(cfg.image_width_inches))

    # Process remaining entries after last image
    if entry_idx < len(entries):
        entry_idx, added = _flush_speaker_entries(entries, entry_idx, doc, is_json)
        paragraph_count += added

    doc.save(output_docx)
    logger.info(f"Total paragraphs created: {paragraph_count}")
    logger.info(f"Total entries processed: {len(entries)}")
    logger.info(f"Average parts per paragraph: {len(entries) / max(paragraph_count, 1):.2f}")
    logger.info(f"Document saved as {output_docx}")

def _flush_speaker_entries(entries, entry_idx, doc, is_json, stop_at=None):
    """
    Process entries and flush dialogue by speaker, accumulating lines from same speaker.

    Args:
        entries: List of dialogue entries
        entry_idx: Current position in entries list
        doc: Word document to add paragraphs to
        is_json: Whether entries include confidence scores
        stop_at: Stop before the first entry starting at or after this time.
            None flushes everything remaining. A timestamp rather than a
            predicate: the only condition any caller ever wanted was "stop at
            this image's start", and building a closure per image meant one that
            captured the loop variable — harmless while it is called inside the
            same iteration, and silently wrong the day anyone defers the call.

    Returns:
        (new_entry_idx, paragraphs_added)
    """
    pending_speaker = None
    pending_dialogue_parts = []
    paragraphs_added = 0

    while entry_idx < len(entries):
        entry = entries[entry_idx]

        # `is not None`, not truthiness: 0.0 is a legitimate stop time.
        if stop_at is not None and entry['start'] >= stop_at:
            break

        speaker = entry['speaker']
        dialogue = entry['text']
        confidence = entry['confidence']

        # Flush when speaker changes
        if pending_speaker is not None and speaker != pending_speaker:
            _add_dialogue_paragraph(doc, pending_speaker, pending_dialogue_parts, is_json)
            paragraphs_added += 1
            pending_speaker = None
            pending_dialogue_parts = []

        if pending_speaker is None:
            pending_speaker = speaker

        pending_dialogue_parts.append({'text': dialogue, 'confidence': confidence})
        entry_idx += 1

    # Flush remaining speaker
    if pending_speaker is not None:
        _add_dialogue_paragraph(doc, pending_speaker, pending_dialogue_parts, is_json)
        paragraphs_added += 1

    return entry_idx, paragraphs_added


def _add_dialogue_paragraph(doc, speaker, dialogue_parts, is_json):
    """Helper to add a single paragraph with one or more dialogue lines from same speaker."""
    # Add speaker's dialogue
    p = doc.add_paragraph()

    speaker_run = p.add_run(f"{speaker}: ")
    speaker_run.bold = True

    for i, part in enumerate(dialogue_parts):
        text = part['text']
        confidence = part['confidence']

        p.add_run(text)

        if i < len(dialogue_parts) - 1:
            if is_json and confidence < 1.0:
                conf_run = p.add_run(f" ({confidence:.0%}) ")
                conf_run.italic = True
                conf_run.font.size = Pt(9)
            else:
                p.add_run(" ")

    if is_json and dialogue_parts and dialogue_parts[-1]['confidence'] < 1.0:
        conf_run = p.add_run(f" ({dialogue_parts[-1]['confidence']:.0%})")
        conf_run.italic = True
        conf_run.font.size = Pt(9)

def main():
    if len(sys.argv) != 3:
        print("Usage: python3 -m pipeline.merge.storyboard <scene_dir> <session_text>")
        answer = input("use defaults? (y/n): ").strip().lower()
        if answer == 'y':
            scene_dir = Path("./scenes")
            session_text = Path("./session.txt")
        else:
            sys.exit("Usage: python3 -m pipeline.merge.storyboard <scene_dir> <session_text>")
    else:
        scene_dir = Path(sys.argv[1])
        session_text = Path(sys.argv[2])

    csv = list(Path(scene_dir).glob("*.csv"))
    if not scene_dir.exists():
        print(f"Scene directory does not exist: {scene_dir}")
        return
    if not session_text.exists():
        print(f"Session file does not exist: {session_text}")
        return
    if not scene_dir.is_dir():
        sys.exit(f"Scenes is not a directory: {scene_dir}")
    if not csv or not csv[0].exists():
        print(f"CSV not found in Scenes directory: {csv}")
        return

    csv_path = csv[0]
    working_dir = Path.cwd().name
    print(f"Using {scene_dir} images from {csv} against session {session_text} in {working_dir}")
    output_docx = Path.cwd() / (working_dir.replace(" ", "_") + ".docx")

    generate_storyboard(csv_path, session_text, scene_dir, output_docx)

if __name__ == "__main__":
    main()
