#!/usr/bin/env python3
"""
output_compare.py - content comparison of a produced output tree against goldens.

Both integration drivers share this one comparator:

  * run_tests.py runs each stage individually, with output redirected to a
    SEPARATE tree  (tests/integration/output/<key>/cc_output/...).
  * run_orchestrate_tests.py runs the whole orchestrator, output landing BESIDE
    the source            (tests/test_source/<key>/cc_output/...).

The two hierarchies differ only in WHERE cc_output lands; the thing being checked
is identical. So this module compares two directories and is handed the roots by
the caller - it knows nothing about either layout. Point it at the two cc_output
directories and it walks them.

Comparison is by CONTENT, per file type, not by bytes - because several outputs
are legitimately non-deterministic at the byte level:

  * .docx  a zip of XML carrying embedded timestamps and nondeterministic member
           ordering; two runs over the SAME input never match byte-for-byte
           (the "Word docs always differ" case). Compared by extracted text
           (paragraphs + table cells) instead, which is what actually regresses.
  * .json  key order and whitespace are irrelevant; compared as parsed objects.
  * .csv   compared line-by-line after newline normalization - a manual-capture
           CSV can be CRLF on Windows while the golden is LF.
  * everything else (.jpg, .dat, ...) compared by bytes.

Nothing here imports python-docx at module load: the import is deferred into the
.docx comparator so this module still loads in environments (e.g. a CPU-only CI
sandbox) where python-docx is not installed and no .docx is compared.
"""

from __future__ import annotations

import json
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Tuple


# --------------------------------------------------------------------------
# Per-file-type content comparison
# --------------------------------------------------------------------------

def _normalized_lines(text: str) -> List[str]:
    """Split on any newline flavour, so CRLF vs LF is not a difference."""
    return text.replace("\r\n", "\n").replace("\r", "\n").split("\n")


def _json_equal(produced: Path, expected: Path) -> bool:
    a = json.loads(produced.read_text(encoding="utf-8"))
    b = json.loads(expected.read_text(encoding="utf-8"))
    return a == b


def _csv_equal(produced: Path, expected: Path) -> bool:
    return (_normalized_lines(produced.read_text(encoding="utf-8"))
            == _normalized_lines(expected.read_text(encoding="utf-8")))


def _docx_text(path: Path) -> List[str]:
    """Visible text of a .docx: body paragraphs then every table cell.

    python-docx is imported HERE, not at module top, so environments without it
    can still use the rest of the comparator. Interleaving of paragraphs and
    tables is not preserved, but both documents are flattened the same way, so
    identical content still compares equal; a genuine text change still differs.
    Empty/whitespace-only lines are dropped so cosmetic spacing is not a diff.
    """
    from docx import Document  # lazy: only .docx comparisons need python-docx

    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.append(cell.text)
    return [s for s in (ln.strip() for ln in lines) if s]


def _docx_equal(produced: Path, expected: Path) -> bool:
    return _docx_text(produced) == _docx_text(expected)


def _bytes_equal(produced: Path, expected: Path) -> bool:
    return produced.read_bytes() == expected.read_bytes()


# Suffix -> comparator. Anything not listed falls back to a byte compare.
_CONTENT_COMPARATORS = {
    ".json": _json_equal,
    ".csv": _csv_equal,
    ".docx": _docx_equal,
}


def _compare_file(produced: Path, expected: Path) -> Tuple[bool, str]:
    """(equal, detail). detail is non-empty only when the compare itself errored
    (unreadable file, malformed JSON, python-docx missing for a .docx, ...)."""
    comparator = _CONTENT_COMPARATORS.get(expected.suffix.lower(), _bytes_equal)
    try:
        return comparator(produced, expected), ""
    except Exception as e:  # a comparison that cannot run is a failure, not a pass
        return False, f"{type(e).__name__}: {e}"


# --------------------------------------------------------------------------
# Tree comparison
# --------------------------------------------------------------------------

@dataclass
class FileDiff:
    rel: str
    kind: str            # "differ" | "missing" | "extra" | "error" | "no-expected-root"
    detail: str = ""


@dataclass
class TreeComparison:
    produced_root: Path
    expected_root: Path
    matched: List[str] = field(default_factory=list)
    diffs: List[FileDiff] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return not self.diffs

    def as_dict(self) -> dict:
        """JSON-serializable form for test_results.json."""
        return {
            "ok": self.ok,
            "produced_root": str(self.produced_root),
            "expected_root": str(self.expected_root),
            "matched": len(self.matched),
            "diffs": [vars(d) for d in self.diffs],
        }

    def summary_lines(self, limit: int = 25) -> List[str]:
        """Human-readable status, indented for the driver logs."""
        if self.ok:
            return [f"  ✓ output matches expected ({len(self.matched)} files)"]

        labels = {
            "missing": "missing (in expected, not produced)",
            "extra": "extra (produced, not in expected)",
            "differ": "content differs",
            "error": "compare error",
            "no-expected-root": "no expected goldens at",
        }
        lines = [f"  ✗ output differs from expected "
                 f"({len(self.matched)} ok, {len(self.diffs)} problem(s)):"]
        for d in self.diffs[:limit]:
            label = labels.get(d.kind, d.kind)
            tail = f" - {d.detail}" if d.detail else ""
            target = d.rel or str(self.expected_root)
            lines.append(f"      [{label}] {target}{tail}")
        if len(self.diffs) > limit:
            lines.append(f"      ... and {len(self.diffs) - limit} more")
        return lines


def wipe_output_tree(cc_dir, output_root_name: str) -> None:
    """Remove a produced output tree so a run starts from a clean slate.

    Wiping at the START of a run (not the end) is deliberate: a failed run leaves
    its output on disk to inspect, and — crucially for the comparison — every file
    that remains was produced by THIS run. A stage that silently stops producing a
    file then surfaces as `missing`, rather than falsely `matched` against a stale
    copy left by a previous run.

    GUARDED: only ever removes a directory whose leaf name equals
    ``output_root_name`` (the callers pass OUTPUT_ROOT, i.e. "cc_output").
    run_orchestrate_tests.py writes output BESIDE the input media, so a wrong
    target here would delete source recordings — this refuses anything that is not
    the output root. ``output_root_name`` is passed in rather than imported so this
    module keeps no dependency on the pipeline package.
    """
    cc_dir = Path(cc_dir)
    if cc_dir.name != output_root_name:
        raise ValueError(
            f"refusing to wipe {cc_dir!r}: leaf name is not {output_root_name!r}"
        )
    if cc_dir.exists():
        shutil.rmtree(cc_dir)


def _rel_files(root: Path) -> set:
    if not root.exists():
        return set()
    # POSIX-style relative keys so a Windows-produced tree and a golden tree
    # compare on the same string regardless of path separator.
    return {
        str(p.relative_to(root)).replace("\\", "/")
        for p in root.rglob("*")
        if p.is_file()
    }


def compare_output_tree(produced_root, expected_root, allow_extra: bool = False) -> TreeComparison:
    """Compare a produced cc_output tree against its golden.

    Args:
        produced_root: directory the run actually wrote (a cc_output dir).
        expected_root: the golden directory to check against (a cc_output dir).
        allow_extra: when True, files produced but absent from the golden are not
            treated as failures. Default False - an unexpected extra file is
            drift worth surfacing.

    Files are matched by their path relative to each root, then compared by
    content according to file type (see module docstring).
    """
    produced_root = Path(produced_root)
    expected_root = Path(expected_root)
    result = TreeComparison(produced_root, expected_root)

    if not expected_root.exists():
        # No goldens: fail loudly rather than reporting a vacuous pass.
        result.diffs.append(FileDiff("", "no-expected-root", str(expected_root)))
        return result

    produced = _rel_files(produced_root)
    expected = _rel_files(expected_root)

    for rel in sorted(expected - produced):
        result.diffs.append(FileDiff(rel, "missing"))
    if not allow_extra:
        for rel in sorted(produced - expected):
            result.diffs.append(FileDiff(rel, "extra"))
    for rel in sorted(expected & produced):
        equal, detail = _compare_file(produced_root / rel, expected_root / rel)
        if equal:
            result.matched.append(rel)
        else:
            result.diffs.append(FileDiff(rel, "error" if detail else "differ", detail))

    return result
